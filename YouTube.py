"""
YouTube AI Chat Excel Automation Script
=====================================

This script automates the process of reading content from Excel cells, processing it through AI,
and saving the responses back to Excel cells and text files. It's specifically designed for
YouTube content automation workflow.

Features:
- Reads content from Excel cells (C2, C3, C9, C10, C11, C12)
- Processes content through OpenRouter AI API using DeepSeek model
- Saves responses to Excel cells (B4, B6, B9, B10) and text files
- Uses temporary Word files for enhanced content processing (C11, C12)
- Provides user confirmation prompts between steps
- Automatic cleanup of temporary files

Requirements:
- Excel file: D:\\Anant\\Youtube\\ValueProITGyan\\YouTubeVideosList.xlsx
- Sheet: "Shorts_Automation"
- API key in .env file: OPENROUTER_API_KEY

Author: AI Assistant
Date: June 2025
Version: 2.0 (Enhanced with complete Word file content processing)
"""

import requests
import json
import os
from dotenv import load_dotenv
from openpyxl import load_workbook
from docx import Document

# Load environment variables from .env file
load_dotenv()

# Configuration - Excel file path and sheet name
EXCEL_FILE_PATH = r"D:\Anant\Youtube\ValueProITGyan\YouTubeVideosList.xlsx"
SHEET_NAME = "Shorts_Automation"

def get_ai_response(prompt, model="deepseek/deepseek-chat:free"):
    """
    Get AI response from OpenRouter API using the specified model.
    
    Args:
        prompt (str): The prompt to send to the AI model
        model (str): The AI model to use (default: deepseek/deepseek-chat:free)
    
    Returns:
        str: AI response content or None if request fails
    """
    api_key = os.getenv('OPENROUTER_API_KEY', 'your-api-key-here')
    
    # Check if API key is properly configured
    if api_key == 'your-api-key-here':
        print("❌ Error: API key not found. Check your .env file.")
        return None
    
    # Set up headers for OpenRouter API
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://localhost",
        "X-Title": "YouTube AI Chat",
    }
    
    # Prepare request data
    data = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}]
    }
    
    try:
        # Make API request to OpenRouter
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            data=json.dumps(data)
        )
        
        # Process successful response
        if response.status_code == 200:
            result = response.json()
            if 'choices' in result and len(result['choices']) > 0:
                return result['choices'][0]['message']['content']
            else:
                print("❌ No response content found in API response")
                return None
        else:
            # Handle API errors
            print(f"❌ API Error: HTTP {response.status_code}")
            if response.status_code == 401:
                print("🔑 Authentication Error: Invalid API key")
            print("Response:", response.text)
            return None
            
    except Exception as e:
        print(f"❌ Request failed: {e}")
        return None

def read_excel_data():
    """
    Read data from specified Excel cells in the Shorts_Automation sheet.
    
    Uses data_only=True to read actual cell values instead of formulas.
    Reads from cells: C2, C3, C9, C10, C11, C12
    
    Returns:
        dict: Dictionary containing cell names as keys and cell values as values
        None: If file not found or error occurs
    """
    try:
        print(f"📖 Reading Excel file: {EXCEL_FILE_PATH}")
        # Use data_only=True to get cell values instead of formulas
        workbook = load_workbook(EXCEL_FILE_PATH, data_only=True)
        
        # Check if the required sheet exists
        if SHEET_NAME not in workbook.sheetnames:
            print(f"❌ Sheet '{SHEET_NAME}' not found in workbook")
            print(f"Available sheets: {workbook.sheetnames}")
            return None
            
        sheet = workbook[SHEET_NAME]
        # Read data from specified cells (values only, not formulas)
        data = {
            'C2': sheet['C2'].value,
            'C3': sheet['C3'].value,
            'C9': sheet['C9'].value,
            'C10': sheet['C10'].value,
            'C11': sheet['C11'].value,
            'C12': sheet['C12'].value
        }
        
        print("✅ Excel data read successfully (cell values only):")
        for cell, value in data.items():
            preview = str(value)[:50] + "..." if value and len(str(value)) > 50 else str(value)
            print(f"   {cell}: {preview}")
        
        workbook.close()
        return data
        
    except FileNotFoundError:
        print(f"❌ Excel file not found: {EXCEL_FILE_PATH}")
        return None
    except Exception as e:
        print(f"❌ Error reading Excel file: {e}")
        return None

def write_to_excel(cell, value):
    """
    Write data to a specific Excel cell.
    
    Args:
        cell (str): Cell reference (e.g., 'B4', 'B6')
        value (str): Value to write to the cell
    
    Returns:
        bool: True if successful, False if error occurs
    """
    try:
        workbook = load_workbook(EXCEL_FILE_PATH)
        sheet = workbook[SHEET_NAME]
        sheet[cell] = value
        workbook.save(EXCEL_FILE_PATH)
        workbook.close()
        print(f"✅ Written to Excel cell {cell}")
        return True
    except Exception as e:
        print(f"❌ Error writing to Excel: {e}")
        return False

def write_to_text_file(filename, content):
    """
    Write content to a text file in the same directory as the Excel file.
    
    Args:
        filename (str): Base filename without extension
        content (str): Content to write to the file
    
    Returns:
        bool: True if successful, False if error occurs
    """
    try:
        filepath = os.path.join(os.path.dirname(EXCEL_FILE_PATH), f"{filename}.txt")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"✅ Written to text file: {filename}.txt")
        return True
    except Exception as e:
        print(f"❌ Error writing to text file: {e}")
        return False

def create_temp_file(cell_name, content):
    """
    Create temporary text file with cell content for processing.
    
    Args:
        cell_name (str): Name of the cell (e.g., 'C2', 'C3')
        content: Cell content to write to temporary file
    
    Returns:
        str: Filepath of created temporary file or None if error occurs
    """
    try:
        temp_filename = f"temp_{cell_name}.txt"
        temp_filepath = os.path.join(os.path.dirname(EXCEL_FILE_PATH), temp_filename)
        # Ensure we write the actual cell value (not formula) to the temp file
        cell_value = str(content) if content is not None else ""
        with open(temp_filepath, 'w', encoding='utf-8') as f:
            f.write(cell_value)
        print(f"📄 Created temporary file: {temp_filename} (cell value only)")
        return temp_filepath
    except Exception as e:
        print(f"❌ Error creating temporary file: {e}")
        return None

def read_temp_file(filepath):
    """
    Read content from temporary text file.
    
    Args:
        filepath (str): Path to the temporary file
    
    Returns:
        str: File content or None if error occurs
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        print(f"📖 Read content from temporary file")
        return content
    except Exception as e:
        print(f"❌ Error reading temporary file: {e}")
        return None

def cleanup_temp_file(filepath):
    """
    Clean up temporary file after processing.
    
    Args:
        filepath (str): Path to the temporary file to delete
    """
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            print(f"🗑️ Cleaned up temporary file")
    except Exception as e:
        print(f"❌ Error cleaning up temporary file: {e}")

def create_temp_word_file(cell_name, content):
    """
    Create temporary Word file with cell content, preserving formatting and newlines.
    
    This enhanced version splits content by newlines and creates separate paragraphs
    to preserve the original structure and formatting.
    
    Args:
        cell_name (str): Name of the cell (e.g., 'C11', 'C12')
        content: Cell content to write to temporary Word file
    
    Returns:
        str: Filepath of created temporary Word file or None if error occurs
    """
    try:
        temp_filename = f"temp_{cell_name}.docx"
        temp_filepath = os.path.join(os.path.dirname(EXCEL_FILE_PATH), temp_filename)
        
        # Create a new Word document
        doc = Document()
        # Ensure we write the actual cell value (not formula) to the Word file
        cell_value = str(content) if content is not None else ""
        
        # Preserve line breaks and formatting by splitting on newlines
        if '\n' in cell_value:
            # Split content by newlines and add each as separate paragraph
            lines = cell_value.split('\n')
            for line in lines:
                doc.add_paragraph(line)
        else:
            # Single paragraph for content without newlines
            doc.add_paragraph(cell_value)
        
        doc.save(temp_filepath)
        
        print(f"📄 Created temporary Word file: {temp_filename} (preserving all content and formatting)")
        return temp_filepath
    except Exception as e:
        print(f"❌ Error creating temporary Word file: {e}")
        return None

def read_temp_word_file(filepath):
    """
    Read complete content from temporary Word file preserving all formatting and newlines.
    
    This enhanced version preserves the original document structure by reading all
    paragraphs and tables, maintaining newlines and formatting.
    
    Args:
        filepath (str): Path to the temporary Word file
    
    Returns:
        str: Complete file content with preserved formatting or None if error occurs
    """
    try:
        doc = Document(filepath)
        content = ""
        
        # Read all paragraphs and preserve original structure
        for i, paragraph in enumerate(doc.paragraphs):
            content += paragraph.text
            # Add newline after each paragraph except the last one
            if i < len(doc.paragraphs) - 1:
                content += "\n"
        
        # Also check for any tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += "\n" + cell.text
        
        print(f"📖 Read complete content from temporary Word file ({len(content)} characters)")
        print(f"📝 Content preview (first 200 chars): {content[:200]}...")
        
        # Return the complete content without stripping to preserve all formatting
        return content
    except Exception as e:
        print(f"❌ Error reading temporary Word file: {e}")
        return None

def get_user_confirmation(message):
    """
    Get user confirmation before proceeding with the next step.
    
    Args:
        message (str): Message to display to the user
    
    Returns:
        bool: True if user confirms (presses 'Y'), False otherwise
    """
    while True:
        response = input(f"\n{message}\nPress 'Y' to continue, any other key to stop: ").strip()
        if response.upper() == 'Y':
            return True
        else:
            return False

def main():
    """
    Main automation workflow that processes Excel content through AI.
    
    Workflow Steps:
    1. Read Excel data from specified cells (C2, C3, C9, C10, C11, C12)
    2. Process C2 → B4 (Excel cell via temporary text file)
    3. Process C3 → B6 (Excel cell via temporary text file)
    4. Process C9 → B9 (Excel cell via temporary text file)
    5. Process C10 → B10 (Excel cell via temporary text file)
    6. Process C11 → ShortEng_AT.txt (Text file via temporary Word file)
    7. Process C12 → ShortHindi_AT.txt (Text file via temporary Word file)
    
    Each step includes:
    - User confirmation prompt
    - Temporary file creation
    - AI processing with enhanced prompts
    - Response saving    - File cleanup
    """
    print("🚀 Starting YouTube AI Chat Excel Automation")
    print("=" * 60)
    
    # ==========================================
    # STEP 1: Initialize - Read Excel Data
    # ==========================================
    excel_data = read_excel_data()
    if not excel_data:
        print("❌ Failed to read Excel data. Exiting.")
        return
    
    print("\n📋 Excel data loaded successfully!")
    
    # ==========================================
    # STEP 2: Process C2 → B4 (Excel Cell)
    # ==========================================
    if excel_data['C2']:
        print(f"\n🔄 STEP 1: Processing C2 content")
        print(f"📝 Content preview: {str(excel_data['C2'])[:100]}...")
        
        # Create temporary file with C2 content
        temp_file = create_temp_file("C2", excel_data['C2'])
        if temp_file:
            # Read content from temporary file
            temp_content = read_temp_file(temp_file)
            if temp_content:
                # Generate AI response
                prompt_c2 = f"Based on this content from the temporary file: '{temp_content}', generate an appropriate response for cell B4."
                response_b4 = get_ai_response(prompt_c2)
                
                if response_b4:
                    print(f"\n📝 AI Response for B4:\n{response_b4}")
                    # Save to Excel B4
                    write_to_excel('B4', response_b4)
                    print("✅ Step 1 completed: C2 → B4")
                else:
                    print("❌ Failed to get AI response for B4")
            
            # Cleanup temporary file
            cleanup_temp_file(temp_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 1 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return    
    # ==========================================
    # STEP 3: Process C3 → B6 (Excel Cell)
    # ==========================================
    if excel_data['C3']:
        print(f"\n🔄 STEP 2: Processing C3 content")
        print(f"📝 Content preview: {str(excel_data['C3'])[:100]}...")
        
        # Create temporary file with C3 content
        temp_file = create_temp_file("C3", excel_data['C3'])
        if temp_file:
            # Read content from temporary file
            temp_content = read_temp_file(temp_file)
            if temp_content:
                # Generate AI response
                prompt_c3 = f"Based on this content from the temporary file: '{temp_content}', generate an appropriate response for cell B6."
                response_b6 = get_ai_response(prompt_c3)
                
                if response_b6:
                    print(f"\n📝 AI Response for B6:\n{response_b6}")
                    # Save to Excel B6
                    write_to_excel('B6', response_b6)
                    print("✅ Step 2 completed: C3 → B6")
                else:
                    print("❌ Failed to get AI response for B6")
            
            # Cleanup temporary file
            cleanup_temp_file(temp_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 2 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return    
    # ==========================================
    # STEP 4: Process C9 → B9 (Excel Cell)
    # ==========================================
    if excel_data['C9']:
        print(f"\n🔄 STEP 3: Processing C9 content")
        print(f"📝 Content preview: {str(excel_data['C9'])[:100]}...")
        
        # Create temporary file with C9 content
        temp_file = create_temp_file("C9", excel_data['C9'])
        if temp_file:
            # Read content from temporary file
            temp_content = read_temp_file(temp_file)
            if temp_content:
                # Generate AI response
                prompt_c9 = f"Based on this content from the temporary file: '{temp_content}', generate an appropriate response for cell B9."
                response_b9 = get_ai_response(prompt_c9)
                
                if response_b9:
                    print(f"\n📝 AI Response for B9:\n{response_b9}")
                    # Save to Excel B9
                    write_to_excel('B9', response_b9)
                    print("✅ Step 3 completed: C9 → B9")
                else:
                    print("❌ Failed to get AI response for B9")
            
            # Cleanup temporary file
            cleanup_temp_file(temp_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 3 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return    
    # ==========================================
    # STEP 5: Process C10 → B10 (Excel Cell)
    # ==========================================
    if excel_data['C10']:
        print(f"\n🔄 STEP 4: Processing C10 content")
        print(f"📝 Content preview: {str(excel_data['C10'])[:100]}...")
        
        # Create temporary file with C10 content
        temp_file = create_temp_file("C10", excel_data['C10'])
        if temp_file:
            # Read content from temporary file
            temp_content = read_temp_file(temp_file)
            if temp_content:
                # Generate AI response
                prompt_c10 = f"Based on this content from the temporary file: '{temp_content}', generate an appropriate response for cell B10."
                response_b10 = get_ai_response(prompt_c10)
                
                if response_b10:
                    print(f"\n📝 AI Response for B10:\n{response_b10}")
                    # Save to Excel B10
                    write_to_excel('B10', response_b10)
                    print("✅ Step 4 completed: C10 → B10")
                else:
                    print("❌ Failed to get AI response for B10")
            
            # Cleanup temporary file
            cleanup_temp_file(temp_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 4 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return    
    # ==========================================
    # STEP 6: Process C11 → ShortEng_AT.txt (Text File via Word)
    # ==========================================
    if excel_data['C11']:
        print(f"\n🔄 STEP 5: Processing C11 content")
        print(f"📝 Content preview: {str(excel_data['C11'])[:100]}...")
        
        # Create temporary Word file with C11 content
        temp_word_file = create_temp_word_file("C11", excel_data['C11'])
        if temp_word_file:            # Read content from temporary Word file
            temp_content = read_temp_word_file(temp_word_file)
            if temp_content:
                # Generate AI response with complete content
                prompt_c11 = f"""COMPLETE CONTENT from Word file (preserving all formatting and newlines):

{temp_content}

Based on the COMPLETE content above from the temporary Word file, generate an appropriate English short response for ShortEng_AT file. Please process the entire content including all lines, paragraphs, and formatting."""
                response_c11 = get_ai_response(prompt_c11)
                
                if response_c11:
                    print(f"\n📝 AI Response for ShortEng_AT:\n{response_c11}")
                    # Save to ShortEng_AT text file
                    write_to_text_file('ShortEng_AT', response_c11)
                    print("✅ Step 5 completed: C11 → ShortEng_AT.txt (via temp Word file)")
                else:
                    print("❌ Failed to get AI response for ShortEng_AT")
            
            # Cleanup temporary Word file
            cleanup_temp_file(temp_word_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 5 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return    
    # ==========================================
    # STEP 7: Process C12 → ShortHindi_AT.txt (Text File via Word)
    # ==========================================
    if excel_data['C12']:
        print(f"\n🔄 STEP 6: Processing C12 content")
        print(f"📝 Content preview: {str(excel_data['C12'])[:100]}...")
        
        # Create temporary Word file with C12 content
        temp_word_file = create_temp_word_file("C12", excel_data['C12'])
        if temp_word_file:            # Read content from temporary Word file
            temp_content = read_temp_word_file(temp_word_file)
            if temp_content:
                # Generate AI response with complete content
                prompt_c12 = f"""COMPLETE CONTENT from Word file (preserving all formatting and newlines):

{temp_content}

Based on the COMPLETE content above from the temporary Word file, generate an appropriate Hindi short response for ShortHindi_AT file. Please process the entire content including all lines, paragraphs, and formatting."""
                response_c12 = get_ai_response(prompt_c12)
                
                if response_c12:
                    print(f"\n📝 AI Response for ShortHindi_AT:\n{response_c12}")
                    # Save to ShortHindi_AT text file
                    write_to_text_file('ShortHindi_AT', response_c12)
                    print("✅ Step 6 completed: C12 → ShortHindi_AT.txt (via temp Word file)")
                else:
                    print("❌ Failed to get AI response for ShortHindi_AT")
            
            # Cleanup temporary Word file
            cleanup_temp_file(temp_word_file)    
    # ==========================================
    # WORKFLOW COMPLETION SUMMARY
    # ==========================================
    print("\n🎉 Excel automation workflow completed successfully!")
    print("=" * 60)
    print("📋 Summary of completed actions:")
    print("   ✅ Step 1: C2 → B4 (Excel cell)")
    print("   ✅ Step 2: C3 → B6 (Excel cell)")
    print("   ✅ Step 3: C9 → B9 (Excel cell)")
    print("   ✅ Step 4: C10 → B10 (Excel cell)")
    print("   ✅ Step 5: C11 → ShortEng_AT.txt (Text file via temp Word file)")
    print("   ✅ Step 6: C12 → ShortHindi_AT.txt (Text file via temp Word file)")
    print("=" * 60)

# ==========================================
# SCRIPT EXECUTION ENTRY POINT
# ==========================================
if __name__ == "__main__":
    main()
