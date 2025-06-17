# Test script for Word file functionality
try:
    from docx import Document
    print("✅ python-docx library is available")
    
    # Test creating a Word document
    doc = Document()
    doc.add_paragraph("Test content")
    doc.save("test_doc.docx")
    print("✅ Word document creation test passed")
    
    # Test reading a Word document
    doc2 = Document("test_doc.docx")
    content = ""
    for paragraph in doc2.paragraphs:
        content += paragraph.text
    print(f"✅ Word document reading test passed: '{content}'")
    
    # Clean up
    import os
    os.remove("test_doc.docx")
    print("✅ Test completed successfully")
    
except ImportError:
    print("❌ python-docx library not installed")
    print("💡 Run: pip install python-docx")
except Exception as e:
    print(f"❌ Error: {e}")
