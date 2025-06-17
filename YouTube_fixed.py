import requests
import json
import os
from dotenv import load_dotenv
from openpyxl import load_workbook
from docx import Document

# Load environment variables from .env file
load_dotenv()

# Configuration
EXCEL_FILE_PATH = r"D:\Anant\Youtube\ValueProITGyan\YouTubeVideosList.xlsx"
SHEET_NAME = "Shorts_Automation"

def get_ai_response(prompt, model="deepseek/deepseek-chat:free"):
    """Get AI response from OpenRouter API"""
    api_key = os.getenv('OPENROUTER_API_KEY', 'your-api-key-here')
    
    if api_key == 'your-api-key-here':
        print("❌ Error: API key not found. Check your .env file.")
        return None
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://localhost",
        "X-Title": "YouTube AI Chat",
    }
    
    data = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}]
    }
    
    try:
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            data=json.dumps(data)
        )
        
        if response.status_code == 200:
            result = response.json()
            if 'choices' in result and len(result['choices']) > 0:
                return result['choices'][0]['message']['content']
            else:
                print("❌ No response content found in API response")
                return None
        else:
            print(f"❌ API Error: HTTP {response.status_code}")
            if response.status_code == 401:
                print("🔑 Authentication Error: Invalid API key")
            print("Response:", response.text)
            return None
            
    except Exception as e:
        print(f"❌ Request failed: {e}")
        return None

def read_excel_data():
    """Read data from Excel file"""
    try:
        print(f"📖 Reading Excel file: {EXCEL_FILE_PATH}")
        # Use data_only=True to get cell values instead of formulas
        workbook = load_workbook(EXCEL_FILE_PATH, data_only=True)
        
        if SHEET_NAME not in workbook.sheetnames:
            print(f"❌ Sheet '{SHEET_NAME}' not found in workbook")
            print(f"Available sheets: {workbook.sheetnames}")
            return None
            
        sheet = workbook[SHEET_NAME]
        # Read data from specified cells (values only, not formulas)
        data = {
            'C2': sheet['C2'].value,
            'C3': sheet['C3'].value,
            'C9': sheet['C9'].value,
            'C10': sheet['C10'].value,
            'C11': sheet['C11'].value,
            'C12': sheet['C12'].value
        }
        
        print("✅ Excel data read successfully (cell values only):")
        for cell, value in data.items():
            print(f"   {cell}: {value}")
        
        workbook.close()
        return data
        
    except FileNotFoundError:
        print(f"❌ Excel file not found: {EXCEL_FILE_PATH}")
        return None
    except Exception as e:
        print(f"❌ Error reading Excel file: {e}")
        return None

def write_to_excel(cell, value):
    """Write data to Excel cell"""
    try:
        workbook = load_workbook(EXCEL_FILE_PATH)
        sheet = workbook[SHEET_NAME]
        sheet[cell] = value
        workbook.save(EXCEL_FILE_PATH)
        workbook.close()
        print(f"✅ Written to Excel cell {cell}")
        return True
    except Exception as e:
        print(f"❌ Error writing to Excel: {e}")
        return False

def write_to_text_file(filename, content):
    """Write content to text file"""
    try:
        filepath = os.path.join(os.path.dirname(EXCEL_FILE_PATH), f"{filename}.txt")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"✅ Written to text file: {filename}.txt")
        return True
    except Exception as e:
        print(f"❌ Error writing to text file: {e}")
        return False

def create_temp_file(cell_name, content):
    """Create temporary text file with cell content (values only, not formulas)"""
    try:
        temp_filename = f"temp_{cell_name}.txt"
        temp_filepath = os.path.join(os.path.dirname(EXCEL_FILE_PATH), temp_filename)
        # Ensure we write the actual cell value (not formula) to the temp file
        cell_value = str(content) if content is not None else ""
        with open(temp_filepath, 'w', encoding='utf-8') as f:
            f.write(cell_value)
        print(f"📄 Created temporary file: {temp_filename} (cell value only)")
        return temp_filepath
    except Exception as e:
        print(f"❌ Error creating temporary file: {e}")
        return None

def read_temp_file(filepath):
    """Read content from temporary file"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        print(f"📖 Read content from temporary file")
        return content
    except Exception as e:
        print(f"❌ Error reading temporary file: {e}")
        return None

def cleanup_temp_file(filepath):
    """Clean up temporary file"""
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            print(f"🗑️ Cleaned up temporary file")
    except Exception as e:
        print(f"❌ Error cleaning up temporary file: {e}")

def create_temp_word_file(cell_name, content):
    """Create temporary Word file with cell content (values only, not formulas)"""
    try:
        temp_filename = f"temp_{cell_name}.docx"
        temp_filepath = os.path.join(os.path.dirname(EXCEL_FILE_PATH), temp_filename)
        
        # Create a new Word document
        doc = Document()
        # Ensure we write the actual cell value (not formula) to the Word file
        cell_value = str(content) if content is not None else ""
        doc.add_paragraph(cell_value)
        doc.save(temp_filepath)
        
        print(f"📄 Created temporary Word file: {temp_filename} (cell value only)")
        return temp_filepath
    except Exception as e:
        print(f"❌ Error creating temporary Word file: {e}")
        return None

def read_temp_word_file(filepath):
    """Read content from temporary Word file"""
    try:
        doc = Document(filepath)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        print(f"📖 Read content from temporary Word file")
        return content.strip()
    except Exception as e:
        print(f"❌ Error reading temporary Word file: {e}")
        return None

def get_user_confirmation(message):
    """Get user confirmation before proceeding"""
    while True:
        response = input(f"\n{message}\nPress 'Y' to continue, any other key to stop: ").strip()
        if response.upper() == 'Y':
            return True
        else:
            return False

def main():
    """Main automation workflow"""
    print("🚀 Starting YouTube AI Chat Excel Automation")
    print("=" * 60)
    
    # Step 1: Read Excel data
    excel_data = read_excel_data()
    if not excel_data:
        print("❌ Failed to read Excel data. Exiting.")
        return
    
    print("\n📋 Excel data loaded successfully!")
    
    # Step 2: Process C2 content for B4
    if excel_data['C2']:
        print(f"\n🔄 STEP 1: Processing C2 content")
        print(f"📝 Content preview: {str(excel_data['C2'])[:100]}...")
        
        # Create temporary file with C2 content
        temp_file = create_temp_file("C2", excel_data['C2'])
        if temp_file:
            # Read content from temporary file
            temp_content = read_temp_file(temp_file)
            if temp_content:
                # Generate AI response
                prompt_c2 = f"Based on this content from the temporary file: '{temp_content}', generate an appropriate response for cell B4."
                response_b4 = get_ai_response(prompt_c2)
                
                if response_b4:
                    print(f"\n📝 AI Response for B4:\n{response_b4}")
                    # Save to Excel B4
                    write_to_excel('B4', response_b4)
                    print("✅ Step 1 completed: C2 → B4")
                else:
                    print("❌ Failed to get AI response for B4")
            
            # Cleanup temporary file
            cleanup_temp_file(temp_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 1 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return
    
    # Step 3: Process C3 content for B6
    if excel_data['C3']:
        print(f"\n🔄 STEP 2: Processing C3 content")
        print(f"📝 Content preview: {str(excel_data['C3'])[:100]}...")
        
        # Create temporary file with C3 content
        temp_file = create_temp_file("C3", excel_data['C3'])
        if temp_file:
            # Read content from temporary file
            temp_content = read_temp_file(temp_file)
            if temp_content:
                # Generate AI response
                prompt_c3 = f"Based on this content from the temporary file: '{temp_content}', generate an appropriate response for cell B6."
                response_b6 = get_ai_response(prompt_c3)
                
                if response_b6:
                    print(f"\n📝 AI Response for B6:\n{response_b6}")
                    # Save to Excel B6
                    write_to_excel('B6', response_b6)
                    print("✅ Step 2 completed: C3 → B6")
                else:
                    print("❌ Failed to get AI response for B6")
            
            # Cleanup temporary file
            cleanup_temp_file(temp_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 2 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return
    
    # Step 4: Process C9 content for B9
    if excel_data['C9']:
        print(f"\n🔄 STEP 3: Processing C9 content")
        print(f"📝 Content preview: {str(excel_data['C9'])[:100]}...")
        
        # Create temporary file with C9 content
        temp_file = create_temp_file("C9", excel_data['C9'])
        if temp_file:
            # Read content from temporary file
            temp_content = read_temp_file(temp_file)
            if temp_content:
                # Generate AI response
                prompt_c9 = f"Based on this content from the temporary file: '{temp_content}', generate an appropriate response for cell B9."
                response_b9 = get_ai_response(prompt_c9)
                
                if response_b9:
                    print(f"\n📝 AI Response for B9:\n{response_b9}")
                    # Save to Excel B9
                    write_to_excel('B9', response_b9)
                    print("✅ Step 3 completed: C9 → B9")
                else:
                    print("❌ Failed to get AI response for B9")
            
            # Cleanup temporary file
            cleanup_temp_file(temp_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 3 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return
    
    # Step 5: Process C10 content for B10
    if excel_data['C10']:
        print(f"\n🔄 STEP 4: Processing C10 content")
        print(f"📝 Content preview: {str(excel_data['C10'])[:100]}...")
        
        # Create temporary file with C10 content
        temp_file = create_temp_file("C10", excel_data['C10'])
        if temp_file:
            # Read content from temporary file
            temp_content = read_temp_file(temp_file)
            if temp_content:
                # Generate AI response
                prompt_c10 = f"Based on this content from the temporary file: '{temp_content}', generate an appropriate response for cell B10."
                response_b10 = get_ai_response(prompt_c10)
                
                if response_b10:
                    print(f"\n📝 AI Response for B10:\n{response_b10}")
                    # Save to Excel B10
                    write_to_excel('B10', response_b10)
                    print("✅ Step 4 completed: C10 → B10")
                else:
                    print("❌ Failed to get AI response for B10")
            
            # Cleanup temporary file
            cleanup_temp_file(temp_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 4 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return
    
    # Step 6: Process C11 content for ShortEng_AT text file (using temp Word file)
    if excel_data['C11']:
        print(f"\n🔄 STEP 5: Processing C11 content")
        print(f"📝 Content preview: {str(excel_data['C11'])[:100]}...")
        
        # Create temporary Word file with C11 content
        temp_word_file = create_temp_word_file("C11", excel_data['C11'])
        if temp_word_file:
            # Read content from temporary Word file
            temp_content = read_temp_word_file(temp_word_file)
            if temp_content:
                # Generate AI response
                prompt_c11 = f"Based on this content from the temporary Word file: '{temp_content}', generate an appropriate English short response for ShortEng_AT file."
                response_c11 = get_ai_response(prompt_c11)
                
                if response_c11:
                    print(f"\n📝 AI Response for ShortEng_AT:\n{response_c11}")
                    # Save to ShortEng_AT text file
                    write_to_text_file('ShortEng_AT', response_c11)
                    print("✅ Step 5 completed: C11 → ShortEng_AT.txt (via temp Word file)")
                else:
                    print("❌ Failed to get AI response for ShortEng_AT")
            
            # Cleanup temporary Word file
            cleanup_temp_file(temp_word_file)
        
        # User confirmation to continue
        if not get_user_confirmation("🤔 Step 5 completed. Continue with script execution?"):
            print("⏹️ Script execution stopped by user.")
            return
    
    # Step 7: Process C12 content for ShortHindi_AT text file (using temp Word file)
    if excel_data['C12']:
        print(f"\n🔄 STEP 6: Processing C12 content")
        print(f"📝 Content preview: {str(excel_data['C12'])[:100]}...")
        
        # Create temporary Word file with C12 content
        temp_word_file = create_temp_word_file("C12", excel_data['C12'])
        if temp_word_file:
            # Read content from temporary Word file
            temp_content = read_temp_word_file(temp_word_file)
            if temp_content:
                # Generate AI response
                prompt_c12 = f"Based on this content from the temporary Word file: '{temp_content}', generate an appropriate Hindi short response for ShortHindi_AT file."
                response_c12 = get_ai_response(prompt_c12)
                
                if response_c12:
                    print(f"\n📝 AI Response for ShortHindi_AT:\n{response_c12}")
                    # Save to ShortHindi_AT text file
                    write_to_text_file('ShortHindi_AT', response_c12)
                    print("✅ Step 6 completed: C12 → ShortHindi_AT.txt (via temp Word file)")
                else:
                    print("❌ Failed to get AI response for ShortHindi_AT")
            
            # Cleanup temporary Word file
            cleanup_temp_file(temp_word_file)
    
    print("\n🎉 Excel automation workflow completed successfully!")
    print("=" * 60)
    print("📋 Summary of completed actions:")
    print("   ✅ Step 1: C2 → B4 (Excel cell)")
    print("   ✅ Step 2: C3 → B6 (Excel cell)")
    print("   ✅ Step 3: C9 → B9 (Excel cell)")
    print("   ✅ Step 4: C10 → B10 (Excel cell)")
    print("   ✅ Step 5: C11 → ShortEng_AT.txt (Text file via temp Word file)")
    print("   ✅ Step 6: C12 → ShortHindi_AT.txt (Text file via temp Word file)")
    print("=" * 60)

if __name__ == "__main__":
    main()
