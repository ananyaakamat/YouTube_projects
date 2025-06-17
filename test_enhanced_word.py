import os
from docx import Document

def test_word_file_processing():
    """Test the enhanced Word file processing functions"""
    
    # Test content with multiple lines and formatting
    test_content = """This is line 1 of test content.

This is line 3 after a blank line.
Here's line 4 with some special characters: 🚀📝✅

Line 6 with more content and formatting.
Final line of test content."""
    
    print("🧪 Testing enhanced Word file processing...")
    print(f"📝 Original content:\n{repr(test_content)}")
    print(f"📝 Original content length: {len(test_content)} characters")
    
    # Create temporary Word file (simulating the enhanced function)
    temp_filename = "test_temp_word.docx"
    temp_filepath = os.path.join(os.getcwd(), temp_filename)
    
    try:
        # Enhanced Word file creation
        doc = Document()
        if '\n' in test_content:
            lines = test_content.split('\n')
            for line in lines:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(test_content)
        
        doc.save(temp_filepath)
        print(f"✅ Created test Word file: {temp_filename}")
        
        # Enhanced Word file reading
        doc = Document(temp_filepath)
        content = ""
        
        for i, paragraph in enumerate(doc.paragraphs):
            content += paragraph.text
            if i < len(doc.paragraphs) - 1:
                content += "\n"
        
        print(f"📖 Read content from Word file:")
        print(f"📝 Read content:\n{repr(content)}")
        print(f"📝 Read content length: {len(content)} characters")
        
        # Compare original vs read content
        if content == test_content:
            print("✅ SUCCESS: Content matches exactly!")
        else:
            print("❌ DIFFERENCE: Content doesn't match")
            print(f"Original: {len(test_content)} chars")
            print(f"Read:     {len(content)} chars")
        
        # Cleanup
        if os.path.exists(temp_filepath):
            os.remove(temp_filepath)
            print("🗑️ Cleaned up test file")
            
    except Exception as e:
        print(f"❌ Error in test: {e}")
        
    print("🎯 Test completed!")

if __name__ == "__main__":
    test_word_file_processing()
